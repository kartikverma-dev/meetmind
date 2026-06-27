"""Meeting upload and retrieval routes protected by JWT auth with mock fallback, magic bytes, and path sanitization."""

import logging
import os
import tempfile
import uuid
import asyncio
import pathlib
import json
import subprocess
from datetime import datetime, timezone
from pathlib import Path
from uuid import UUID

from fastapi import APIRouter, File, Form, HTTPException, UploadFile, Depends, Request
from fastapi.responses import StreamingResponse

from middleware.rate_limit import limiter, LIMIT_UPLOAD
from utils.sanitize import sanitize_filename, sanitize_text
from utils.magic_bytes import verify_file_type
from config import get_settings
from services.supabase_client import SupabaseNotConfiguredError, get_supabase
from models.schemas import (
    MeetingResponse,
    MeetingStatus,
    MeetingUploadResponse,
    MOM,
)
from services.ai_processor import process_transcript, get_meeting_title
from services.transcriber import transcribe_audio
from routes.auth import get_current_user, check_limits

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/meetings", tags=["meetings"])

# Allowed audio/video extensions for upload
ALLOWED_EXTENSIONS = {
    ".mp3", ".mp4", ".wav", ".m4a", ".ogg", ".webm", ".flac", ".mkv", ".avi",
}

# In-memory mock database for mock mode
MOCK_MEETINGS = [
    {
        "id": "e402b874-9c8a-4933-8551-789a456123cd",
        "user_id": "11111111-1111-1111-1111-111111111111",
        "title": "Q3 Planning & Execution Sync",
        "transcript": "Let's review the Q3 targets. We need to ship the summary module by next week. Alice is leading it. Bob is writing test cases. We will launch on Tuesday.",
        "mom": {
            "attendees": ["Alice", "Bob", "Charlie"],
            "date": "2026-06-25",
            "agenda": ["Q3 timeline alignment", "Summary module ship date"],
            "decisions": ["Ship summary module by next week", "Launch on Tuesday"],
            "action_items": [
                {"task": "Ship summary module", "owner": "Alice", "deadline": "2026-06-30"},
                {"task": "Write test cases", "owner": "Bob", "deadline": "2026-06-29"}
            ]
        },
        "summary": "- Alice will complete the summary module shipping work.\n- Bob is tasked with writing all related test cases by Monday.\n- The official release is scheduled for Tuesday morning.\n- Charlie will coordinate the Vercel/Render deployments.\n- Team agreed to sync again on Friday.",
        "created_at": "2026-06-25T09:00:00+00:00",
        "status": "done"
    }
]

MAX_FILE_SIZE_FREE = 100 * 1024 * 1024   # 100MB
MAX_FILE_SIZE_PRO  = 500 * 1024 * 1024   # 500MB
MAX_DURATION_FREE  = 30 * 60             # 30 minutes in seconds
MAX_DURATION_PRO   = 3 * 60 * 60         # 3 hours in seconds

async def run_db_query(func, *args, **kwargs):
    """Run a synchronous Supabase query in a separate thread with a 10s timeout."""
    try:
        return await asyncio.wait_for(
            asyncio.to_thread(func, *args, **kwargs),
            timeout=10.0
        )
    except asyncio.TimeoutError as e:
        logger.error("Supabase database query timed out: %s", e)
        raise HTTPException(
            status_code=408,
            detail="Database request timed out."
        ) from e

def _validate_extension(filename: str) -> str:
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type. Allowed: mp3, mp4, wav, m4a, ogg, webm, flac, mkv, avi",
        )
    return ext


from pydantic import BaseModel

# Models for additional meeting routes
class PasteTranscriptRequest(BaseModel):
    title: str
    transcript: str

class UpdateMeetingRequest(BaseModel):
    title: str

class PublicToggleRequest(BaseModel):
    is_public: bool

class SearchResultResponse(BaseModel):
    id: UUID
    title: str
    snippet: str
    created_at: str


@router.post("/upload", response_model=MeetingUploadResponse)
@limiter.limit(LIMIT_UPLOAD)
async def upload_meeting(
    request: Request,
    file: UploadFile = File(...),
    title: str = Form(default="Untitled Meeting"),
    language: str = Form(default="auto"),
    current_user = Depends(get_current_user),
    profile = Depends(check_limits),
):
    """
    Upload an audio/video file, transcribe with Whisper, process with Gemini,
    and persist the meeting to Supabase.
    """
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    ext = _validate_extension(file.filename)
    resolved_user_id = str(current_user.id)
    
    # Sanitize title
    title = sanitize_text(title, max_length=100)

    # Read file content safely
    content = await file.read()
    file_size = len(content)

    # BETA MODE: All limits set to PRO tier (500MB)
    is_pro = True
    max_bytes = MAX_FILE_SIZE_PRO
    if file_size > max_bytes:
        raise HTTPException(
            status_code=413,
            detail="File too large. Beta limit is 500MB."
        )

    # Check empty file
    if file_size == 0:
        raise HTTPException(
            status_code=400,
            detail="Uploaded file is empty or corrupt."
        )
    
    # Path Traversal & secure naming configurations
    upload_dir = pathlib.Path(tempfile.gettempdir()) / "meetmind"
    upload_dir.mkdir(parents=True, exist_ok=True)
    
    safe_filename = f"{uuid.uuid4()}{ext}"
    target_path = (upload_dir / safe_filename).resolve()
    
    # Path traversal validation
    if not str(target_path).startswith(str(upload_dir.resolve())):
        raise HTTPException(status_code=400, detail="Invalid path traversal attempt.")

    # Save content to secure temp file
    try:
        with open(target_path, "wb") as f:
            f.write(content)
    except IOError as exc:
        logger.error("Failed to write to temp file: %s", exc)
        raise HTTPException(status_code=500, detail="Internal file saving error.")

    # Magic bytes verification to check actual content type matches extension
    if not verify_file_type(str(target_path), ext):
        if target_path.exists():
            target_path.unlink()
        raise HTTPException(
            status_code=400,
            detail="File signature mismatch. The actual file content does not match its claimed extension."
        )

    settings = get_settings()
    if settings.mock_mode:
        from routes.auth import MOCK_PROFILES
        if resolved_user_id in MOCK_PROFILES:
            MOCK_PROFILES[resolved_user_id]["meetings_used"] += 1

        new_id = str(uuid.uuid4())
        mock_mom_data = {
            "attendees": ["Alice", "Bob", "Demo User"],
            "date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "agenda": ["MeetMind feature overview", "Local Whisper performance verification"],
            "decisions": ["Integrate mock fallback for clean local previewing"],
            "action_items": [
                {"task": "Complete frontend validation tests", "owner": "Demo User", "deadline": "Today", "status": "pending"}
            ]
        }
        new_mock_meeting = {
            "id": new_id,
            "user_id": resolved_user_id,
            "title": title,
            "transcript": f"This is a mock transcribed meeting transcript for '{title}' uploaded in demonstration mode.",
            "mom": mock_mom_data,
            "summary": f"- User successfully uploaded their recording.\n- Local Whisper completed transcribing the audio data.\n- Gemini generated structured Minutes of Meeting (MOM) schema.\n- Executive summary items were extracted.\n- MeetMind successfully stored the meeting details.",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "status": "done",
            "duration": 180,
            "is_public": False,
            "public_slug": None
        }
        MOCK_MEETINGS.append(new_mock_meeting)
        if target_path.exists():
            target_path.unlink()
        return MeetingUploadResponse(
            id=UUID(new_id),
            status=MeetingStatus.DONE,
            message="Meeting processed successfully (Mock Mode)",
        )

    meeting_id = str(uuid.uuid4())
    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        if target_path.exists():
            target_path.unlink()
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    # Increment meetings_used in profiles table
    try:
        profile_res = await run_db_query(
            supabase.table("profiles").select("meetings_used").eq("id", resolved_user_id).maybe_single().execute
        )
        current_used = profile_res.data.get("meetings_used", 0) if (profile_res and profile_res.data) else 0
        await run_db_query(
            supabase.table("profiles").update({"meetings_used": current_used + 1}).eq("id", resolved_user_id).execute
        )
    except Exception as exc:
        logger.error("Failed to increment meetings_used for user %s: %s", resolved_user_id, exc)

    # Create meeting row with 'processing' status before heavy work
    now = datetime.now(timezone.utc).isoformat()
    try:
        await run_db_query(
            supabase.table("meetings").insert(
                {
                    "id": meeting_id,
                    "user_id": resolved_user_id,
                    "title": title,
                    "transcript": None,
                    "mom": None,
                    "summary": None,
                    "created_at": now,
                    "status": MeetingStatus.PROCESSING.value,
                }
            ).execute
        )
    except Exception as exc:
        if target_path.exists():
            target_path.unlink()
        raise HTTPException(status_code=500, detail=f"Failed to create meeting record: {exc}")

    try:
        # Use ffprobe to check audio duration
        try:
            result = subprocess.run(
                ["ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", str(target_path)],
                capture_output=True, text=True, check=True
            )
            ffprobe_data = json.loads(result.stdout)
            duration = float(ffprobe_data["format"]["duration"])
        except Exception as ffprobe_exc:
            logger.error("ffprobe duration check failed: %s", ffprobe_exc)
            raise HTTPException(
                status_code=400,
                detail="Could not read audio file. It may be corrupt."
            )

        max_duration = MAX_DURATION_PRO if is_pro else MAX_DURATION_FREE
        if duration > max_duration:
            limit_msg = "3 hours" if is_pro else "30 minutes"
            raise HTTPException(
                status_code=400,
                detail=f"Audio too long. { 'Pro' if is_pro else 'Free' } tier max is {limit_msg}."
            )

        # Transcribe audio using local Whisper with selected language
        try:
            transcript = await transcribe_audio(str(target_path), language=language)
        except Exception as trans_exc:
            logger.error("Whisper transcription failed: %s", trans_exc)
            raise HTTPException(
                status_code=400,
                detail=f"Invalid or corrupt audio file: {trans_exc}"
            )

        if not transcript or not transcript.strip():
            raise HTTPException(
                status_code=400,
                detail="No speech could be detected in the audio file. Please ensure it is not silent or corrupt."
            )

        # Process transcript with Gemini
        mom, summary = await process_transcript(transcript)

        final_title = title
        if not title or title.strip() == "" or title == "Untitled Meeting":
            try:
                final_title = await get_meeting_title(transcript)
            except Exception as title_exc:
                logger.error("Failed to generate meeting title: %s", title_exc)
                final_title = f"Meeting - {datetime.now(timezone.utc).strftime('%Y-%m-%d')}"

        # Persist results in Supabase
        await run_db_query(
            supabase.table("meetings").update(
                {
                    "title": final_title,
                    "transcript": transcript,
                    "mom": mom.model_dump(),
                    "summary": summary,
                    "status": MeetingStatus.DONE.value,
                    "duration": int(duration),
                }
            ).eq("id", meeting_id).execute
        )

        # Insert action items
        if mom.action_items:
            action_rows = []
            for item in mom.action_items:
                action_rows.append({
                    "meeting_id": meeting_id,
                    "user_id": resolved_user_id,
                    "task": item.task,
                    "owner": item.owner,
                    "deadline": item.deadline,
                    "status": "pending"
                })
            await run_db_query(
                supabase.table("action_items").insert(action_rows).execute
            )

        return MeetingUploadResponse(
            id=UUID(meeting_id),
            status=MeetingStatus.DONE,
            message="Meeting processed successfully",
        )

    except HTTPException:
        # Re-raise explicit HTTPExceptions
        raise
    except Exception as exc:
        logger.exception("Meeting processing failed for %s", meeting_id)
        try:
            await run_db_query(
                supabase.table("meetings").update(
                    {"status": MeetingStatus.FAILED.value}
                ).eq("id", meeting_id).execute
            )
        except Exception as db_exc:
            logger.error("Failed to set meeting status to failed in DB: %s", db_exc)
        raise HTTPException(
            status_code=500,
            detail="Processing failed. Please try again with a valid audio file.",
        )
    finally:
        if target_path.exists():
            try:
                target_path.unlink()
            except Exception as unlink_exc:
                logger.error("Failed to delete temp file: %s", unlink_exc)


@router.get("/search", response_model=list[SearchResultResponse])
async def search_meetings(
    q: str,
    current_user = Depends(get_current_user)
):
    """Search for keywords in meeting titles or transcripts."""
    settings = get_settings()
    resolved_user_id = str(current_user.id)
    if not q or not q.strip():
        raise HTTPException(status_code=400, detail="Search query cannot be empty")
    query = q.strip().lower()

    if settings.mock_mode:
        results = []
        for m in MOCK_MEETINGS:
            if str(m["user_id"]) != resolved_user_id:
                continue
            title_match = query in m.get("title", "").lower()
            transcript_match = query in m.get("transcript", "").lower()
            
            if title_match or transcript_match:
                # Extract snippet
                snippet = ""
                tr = m.get("transcript", "")
                idx = tr.lower().find(query)
                if idx != -1:
                    start = max(0, idx - 60)
                    end = min(len(tr), idx + len(query) + 60)
                    snippet = ("..." if start > 0 else "") + tr[start:end] + ("..." if end < len(tr) else "")
                else:
                    snippet = m.get("summary", "")[:120] + "..."
                    
                results.append(
                    SearchResultResponse(
                        id=UUID(m["id"]),
                        title=m.get("title", "Untitled Meeting"),
                        snippet=snippet,
                        created_at=m.get("created_at")
                    )
                )
        return results

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    try:
        # Search via ILIKE using OR query
        db_res = await run_db_query(
            supabase.table("meetings")
            .select("id, title, transcript, summary, created_at")
            .eq("user_id", resolved_user_id)
            .or_(f"title.ilike.%{query}%,transcript.ilike.%{query}%")
            .execute
        )
        
        results = []
        for row in db_res.data:
            tr = row.get("transcript") or ""
            idx = tr.lower().find(query)
            if idx != -1:
                start = max(0, idx - 60)
                end = min(len(tr), idx + len(query) + 60)
                snippet = ("..." if start > 0 else "") + tr[start:end] + ("..." if end < len(tr) else "")
            else:
                snippet = (row.get("summary") or "")[:120] + "..."
                
            results.append(
                SearchResultResponse(
                    id=UUID(row["id"]),
                    title=row["title"],
                    snippet=snippet,
                    created_at=row["created_at"]
                )
            )
        return results
    except Exception as exc:
        logger.exception("Search failed")
        raise HTTPException(status_code=500, detail=f"Search failed: {exc}")


@router.post("/paste", response_model=MeetingUploadResponse)
async def paste_transcript(
    body: PasteTranscriptRequest,
    current_user = Depends(get_current_user),
    profile = Depends(check_limits)
):
    """Process pasted transcript with Gemini directly."""
    resolved_user_id = str(current_user.id)
    title = sanitize_text(body.title, max_length=100)
    transcript = sanitize_text(body.transcript, max_length=200000)

    if not title or not transcript:
        raise HTTPException(status_code=400, detail="Title and transcript content are required.")

    # Calculate estimated duration (approx 150 words per minute -> ~10 chars per second)
    estimated_duration = max(30, round(len(transcript) / 10))

    settings = get_settings()
    if settings.mock_mode:
        from routes.auth import MOCK_PROFILES
        if resolved_user_id in MOCK_PROFILES:
            MOCK_PROFILES[resolved_user_id]["meetings_used"] += 1

        new_id = str(uuid.uuid4())
        mock_mom_data = {
            "attendees": ["Alice", "Bob", "Demo User"],
            "date": datetime.now(timezone.utc).strftime("%Y-%m-%d"),
            "agenda": ["Pasted transcript sync"],
            "decisions": ["Demonstrate raw text upload"],
            "action_items": [
                {"task": "Verify past dashboard sync", "owner": "Demo User", "deadline": "Today", "status": "pending"}
            ]
        }
        new_mock_meeting = {
            "id": new_id,
            "user_id": resolved_user_id,
            "title": title,
            "transcript": transcript,
            "mom": mock_mom_data,
            "summary": f"- User pasted raw meeting transcript.\n- Gemini processed structure elements.\n- Meeting minutes calculated successfully.",
            "created_at": datetime.now(timezone.utc).isoformat(),
            "status": "done",
            "duration": estimated_duration,
            "is_public": False,
            "public_slug": None
        }
        MOCK_MEETINGS.append(new_mock_meeting)
        return MeetingUploadResponse(
            id=UUID(new_id),
            status=MeetingStatus.DONE,
            message="Meeting processed successfully (Mock Mode)"
        )

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    meeting_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    try:
        # Process transcript with Gemini
        mom, summary = await process_transcript(transcript)
        
        # Save meeting
        await run_db_query(
            supabase.table("meetings").insert(
                {
                    "id": meeting_id,
                    "user_id": resolved_user_id,
                    "title": title,
                    "transcript": transcript,
                    "mom": mom.model_dump(),
                    "summary": summary,
                    "created_at": now,
                    "status": MeetingStatus.DONE.value,
                    "duration": estimated_duration
                }
            ).execute
        )

        # Insert action items
        if mom.action_items:
            action_rows = []
            for item in mom.action_items:
                action_rows.append({
                    "meeting_id": meeting_id,
                    "user_id": resolved_user_id,
                    "task": item.task,
                    "owner": item.owner,
                    "deadline": item.deadline,
                    "status": "pending"
                })
            await run_db_query(
                supabase.table("action_items").insert(action_rows).execute
            )

        # Increment meetings_used count
        profile_res = await run_db_query(
            supabase.table("profiles").select("meetings_used").eq("id", resolved_user_id).maybe_single().execute
        )
        current_used = profile_res.data.get("meetings_used", 0) if (profile_res and profile_res.data) else 0
        await run_db_query(
            supabase.table("profiles").update({"meetings_used": current_used + 1}).eq("id", resolved_user_id).execute
        )

        return MeetingUploadResponse(
            id=UUID(meeting_id),
            status=MeetingStatus.DONE,
            message="Transcript processed successfully"
        )
    except Exception as exc:
        logger.exception("Pasted transcript processing failed")
        raise HTTPException(status_code=500, detail=f"Failed to process pasted transcript: {exc}")


@router.patch("/{meeting_id}", response_model=MeetingResponse)
async def patch_meeting(
    meeting_id: UUID,
    body: UpdateMeetingRequest,
    current_user = Depends(get_current_user)
):
    """Update meeting details (e.g. title)."""
    settings = get_settings()
    resolved_user_id = str(current_user.id)

    if settings.mock_mode:
        for m in MOCK_MEETINGS:
            if str(m["id"]) == str(meeting_id) and str(m["user_id"]) == resolved_user_id:
                m["title"] = body.title
                mom_data = m.get("mom")
                mom = MOM.model_validate(mom_data) if mom_data else None
                return MeetingResponse(
                    id=UUID(m["id"]),
                    user_id=UUID(m["user_id"]),
                    title=m["title"],
                    transcript=m.get("transcript"),
                    mom=mom,
                    summary=m.get("summary"),
                    created_at=m["created_at"],
                    status=MeetingStatus(m["status"]),
                    is_public=m.get("is_public", False),
                    public_slug=m.get("public_slug"),
                    duration=m.get("duration", 0)
                )
        raise HTTPException(status_code=404, detail="Meeting not found")

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    # Check ownership
    result = await run_db_query(
        supabase.table("meetings").select("user_id").eq("id", str(meeting_id)).maybe_single().execute
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Meeting not found")
    if str(result.data["user_id"]) != resolved_user_id:
        raise HTTPException(status_code=403, detail="Permission denied")

    # Update
    update_res = await run_db_query(
        supabase.table("meetings").update({"title": body.title}).eq("id", str(meeting_id)).execute
    )
    row = update_res.data[0]
    mom_data = row.get("mom")
    mom = MOM.model_validate(mom_data) if mom_data else None
    
    return MeetingResponse(
        id=UUID(row["id"]),
        user_id=UUID(row["user_id"]),
        title=row["title"],
        transcript=row.get("transcript"),
        mom=mom,
        summary=row.get("summary"),
        created_at=row["created_at"],
        status=MeetingStatus(row["status"]),
        is_public=row.get("is_public", False),
        public_slug=row.get("public_slug"),
        duration=row.get("duration", 0)
    )


@router.patch("/{meeting_id}/public")
async def toggle_public_meeting(
    meeting_id: UUID,
    body: PublicToggleRequest,
    current_user = Depends(get_current_user)
):
    """Toggle meeting public visibility. Generates public_slug if set to true."""
    settings = get_settings()
    resolved_user_id = str(current_user.id)

    if settings.mock_mode:
        for m in MOCK_MEETINGS:
            if str(m["id"]) == str(meeting_id) and str(m["user_id"]) == resolved_user_id:
                m["is_public"] = body.is_public
                if body.is_public and not m.get("public_slug"):
                    m["public_slug"] = f"mock-slug-{uuid.uuid4().hex[:12]}"
                elif not body.is_public:
                    m["public_slug"] = None
                return {"is_public": m["is_public"], "public_slug": m["public_slug"]}
        raise HTTPException(status_code=404, detail="Meeting not found")

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    # Check ownership
    result = await run_db_query(
        supabase.table("meetings").select("user_id, public_slug").eq("id", str(meeting_id)).maybe_single().execute
    )
    if not result.data:
        raise HTTPException(status_code=404, detail="Meeting not found")
    if str(result.data["user_id"]) != resolved_user_id:
        raise HTTPException(status_code=403, detail="Permission denied")

    slug = result.data.get("public_slug")
    if body.is_public and not slug:
        # Generate new random unique slug
        slug = uuid.uuid4().hex
    elif not body.is_public:
        slug = None

    await run_db_query(
        supabase.table("meetings")
        .update({"is_public": body.is_public, "public_slug": slug})
        .eq("id", str(meeting_id))
        .execute
    )
    
    return {"is_public": body.is_public, "public_slug": slug}


@router.get("", response_model=list[MeetingResponse])
async def list_meetings(
    current_user = Depends(get_current_user),
):
    """List all meetings for the authenticated user, ordered by date."""
    settings = get_settings()
    if settings.mock_mode:
        meetings = []
        for row in MOCK_MEETINGS:
            if str(row["user_id"]) == str(current_user.id):
                mom_data = row.get("mom")
                mom = MOM.model_validate(mom_data) if mom_data else None
                meetings.append(
                    MeetingResponse(
                        id=UUID(row["id"]),
                        user_id=UUID(row["user_id"]),
                        title=row["title"],
                        transcript=row.get("transcript"),
                        mom=mom,
                        summary=row.get("summary"),
                        created_at=row["created_at"],
                        status=MeetingStatus(row["status"]),
                        is_public=row.get("is_public", False),
                        public_slug=row.get("public_slug"),
                        duration=row.get("duration", 0),
                    )
                )
        return meetings

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
    
    result = await run_db_query(
        supabase.table("meetings")
        .select("*")
        .eq("user_id", str(current_user.id))
        .order("created_at", desc=True)
        .execute
    )

    meetings = []
    for row in result.data:
        mom_data = row.get("mom")
        mom = MOM.model_validate(mom_data) if mom_data else None
        meetings.append(
            MeetingResponse(
                id=UUID(row["id"]),
                user_id=UUID(row["user_id"]),
                title=row["title"],
                transcript=row.get("transcript"),
                mom=mom,
                summary=row.get("summary"),
                created_at=row["created_at"],
                status=MeetingStatus(row["status"]),
                is_public=row.get("is_public", False),
                public_slug=row.get("public_slug"),
                duration=row.get("duration", 0),
            )
        )
    return meetings


@router.get("/{meeting_id}", response_model=MeetingResponse)
async def get_meeting(
    meeting_id: UUID,
    current_user = Depends(get_current_user),
):
    """Return full meeting data by ID, verifying user ownership."""
    settings = get_settings()
    if settings.mock_mode:
        row = None
        for m in MOCK_MEETINGS:
            if str(m["id"]) == str(meeting_id):
                row = m
                break
        if not row:
            raise HTTPException(status_code=404, detail="Meeting not found")

        if str(row["user_id"]) != str(current_user.id):
            raise HTTPException(
                status_code=403,
                detail="You do not have permission to access this meeting",
            )

        mom_data = row.get("mom")
        mom = MOM.model_validate(mom_data) if mom_data else None
        return MeetingResponse(
            id=UUID(row["id"]),
            user_id=UUID(row["user_id"]),
            title=row["title"],
            transcript=row.get("transcript"),
            mom=mom,
            summary=row.get("summary"),
            created_at=row["created_at"],
            status=MeetingStatus(row["status"]),
            is_public=row.get("is_public", False),
            public_slug=row.get("public_slug"),
            duration=row.get("duration", 0),
        )

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc
        
    result = await run_db_query(
        supabase.table("meetings")
        .select("*")
        .eq("id", str(meeting_id))
        .maybe_single()
        .execute
    )

    row = result.data
    if not row:
        raise HTTPException(status_code=404, detail="Meeting not found")

    # Enforce meeting ownership check on single retrieval
    if str(row["user_id"]) != str(current_user.id):
        raise HTTPException(
            status_code=403,
            detail="You do not have permission to access this meeting",
        )

    mom_data = row.get("mom")
    mom = MOM.model_validate(mom_data) if mom_data else None

    return MeetingResponse(
        id=UUID(row["id"]),
        user_id=UUID(row["user_id"]),
        title=row["title"],
        transcript=row.get("transcript"),
        mom=mom,
        summary=row.get("summary"),
        created_at=row["created_at"],
        status=MeetingStatus(row["status"]),
        is_public=row.get("is_public", False),
        public_slug=row.get("public_slug"),
        duration=row.get("duration", 0),
    )


@router.get("/{meeting_id}/export/pdf")
async def export_meeting_pdf(
    meeting_id: UUID,
    current_user = Depends(get_current_user),
    profile = Depends(check_limits)
):
    """Export Minutes of Meeting (MOM) as a beautifully formatted PDF."""

    settings = get_settings()
    meeting = None
    
    if settings.mock_mode:
        for m in MOCK_MEETINGS:
            if str(m["id"]) == str(meeting_id):
                meeting = m
                break
    else:
        try:
            supabase = get_supabase()
            result = await run_db_query(
                supabase.table("meetings").select("*").eq("id", str(meeting_id)).maybe_single().execute
            )
            meeting = result.data
        except Exception as exc:
            logger.error("Failed to query meeting: %s", exc)
            raise HTTPException(status_code=500, detail="Database error retrieving meeting")

    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    if str(meeting["user_id"]) != str(current_user.id):
        raise HTTPException(status_code=403, detail="Permission denied")

    mom = meeting.get("mom")
    if not mom:
        raise HTTPException(status_code=400, detail="MOM data not generated yet")

    try:
        from io import BytesIO
        from reportlab.lib.pagesizes import letter
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib import colors

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=54, leftMargin=54, topMargin=54, bottomMargin=54)
        styles = getSampleStyleSheet()

        title_style = ParagraphStyle(
            'DocTitle',
            parent=styles['Heading1'],
            fontName='Helvetica-Bold',
            fontSize=24,
            leading=28,
            textColor=colors.HexColor('#0f172a'),
            spaceAfter=6
        )
        meta_style = ParagraphStyle(
            'DocMeta',
            parent=styles['Normal'],
            fontName='Helvetica-Oblique',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#64748b'),
            spaceAfter=20
        )
        h2_style = ParagraphStyle(
            'DocH2',
            parent=styles['Heading2'],
            fontName='Helvetica-Bold',
            fontSize=14,
            leading=18,
            textColor=colors.HexColor('#1e3a8a'),
            spaceBefore=12,
            spaceAfter=6
        )
        body_style = ParagraphStyle(
            'DocBody',
            parent=styles['Normal'],
            fontName='Helvetica',
            fontSize=10,
            leading=14,
            textColor=colors.HexColor('#334155'),
            spaceAfter=6
        )

        story = []
        title = meeting.get("title", "Untitled Meeting")
        created_at = meeting.get("created_at", "")

        story.append(Paragraph(title, title_style))
        story.append(Paragraph(f"Minutes of Meeting — Generated by MeetMind on {created_at[:10]}", meta_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("Attendees", h2_style))
        attendees = ", ".join(mom.get("attendees", [])) or "None specified"
        story.append(Paragraph(attendees, body_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("Agenda", h2_style))
        for item in mom.get("agenda", []):
            story.append(Paragraph(f"• {item}", body_style))
        if not mom.get("agenda"):
            story.append(Paragraph("None specified", body_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("Decisions Made", h2_style))
        for item in mom.get("decisions", []):
            story.append(Paragraph(f"• {item}", body_style))
        if not mom.get("decisions"):
            story.append(Paragraph("No decisions recorded", body_style))
        story.append(Spacer(1, 10))

        story.append(Paragraph("Action Items", h2_style))
        action_items = mom.get("action_items", [])
        if action_items:
            data = [["Task", "Owner", "Deadline"]]
            for item in action_items:
                data.append([
                    item.get("task", ""),
                    item.get("owner", "Unassigned"),
                    item.get("deadline", "N/A") or "N/A"
                ])

            t = Table(data, colWidths=[280, 110, 110])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#1e3a8a')),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 10),
                ('BOTTOMPADDING', (0,0), (-1,0), 6),
                ('TOPPADDING', (0,0), (-1,0), 6),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.HexColor('#f8fafc'), colors.white]),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
                ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                ('FONTSIZE', (0,1), (-1,-1), 9),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
                ('TOPPADDING', (0,1), (-1,-1), 6),
                ('BOTTOMPADDING', (0,1), (-1,-1), 6),
            ]))
            story.append(t)
        else:
            story.append(Paragraph("No action items recorded", body_style))

        doc.build(story)
        buffer.seek(0)
        
        filename = f"{title.lower().replace(' ', '_')}_mom.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as exc:
        logger.exception("Failed to generate PDF")
        raise HTTPException(status_code=500, detail=f"PDF generation failed: {exc}")


@router.get("/{meeting_id}/export/docx")
async def export_meeting_docx(
    meeting_id: UUID,
    current_user = Depends(get_current_user),
    profile = Depends(check_limits)
):
    """Export Minutes of Meeting (MOM) as a Microsoft Word Document (DOCX)."""

    settings = get_settings()
    meeting = None
    
    if settings.mock_mode:
        for m in MOCK_MEETINGS:
            if str(m["id"]) == str(meeting_id):
                meeting = m
                break
    else:
        try:
            supabase = get_supabase()
            result = await run_db_query(
                supabase.table("meetings").select("*").eq("id", str(meeting_id)).maybe_single().execute
            )
            meeting = result.data
        except Exception as exc:
            logger.error("Failed to query meeting: %s", exc)
            raise HTTPException(status_code=500, detail="Database error retrieving meeting")

    if not meeting:
        raise HTTPException(status_code=404, detail="Meeting not found")

    if str(meeting["user_id"]) != str(current_user.id):
        raise HTTPException(status_code=403, detail="Permission denied")

    mom = meeting.get("mom")
    if not mom:
        raise HTTPException(status_code=400, detail="MOM data not generated yet")

    try:
        from io import BytesIO
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor

        doc = Document()
        
        for section in doc.sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        title = meeting.get("title", "Untitled Meeting")
        created_at = meeting.get("created_at", "")

        t = doc.add_heading(title, level=1)
        t.runs[0].font.name = 'Arial'
        t.runs[0].font.size = Pt(22)
        t.runs[0].font.bold = True
        
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"Minutes of Meeting — Generated by MeetMind on {created_at[:10]}")
        meta_run.font.name = 'Arial'
        meta_run.font.size = Pt(10)
        meta_run.font.italic = True
        meta_run.font.color.rgb = RGBColor(100, 116, 139)

        h = doc.add_heading("Attendees", level=2)
        h.runs[0].font.name = 'Arial'
        h.runs[0].font.size = Pt(14)
        h.runs[0].font.bold = True

        attendees = ", ".join(mom.get("attendees", [])) or "None specified"
        p = doc.add_paragraph(attendees)
        p.runs[0].font.name = 'Arial'
        p.runs[0].font.size = Pt(11)

        h = doc.add_heading("Agenda", level=2)
        h.runs[0].font.name = 'Arial'
        h.runs[0].font.size = Pt(14)
        h.runs[0].font.bold = True

        for item in mom.get("agenda", []):
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(11)
        if not mom.get("agenda"):
            p = doc.add_paragraph("None specified")
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(11)

        h = doc.add_heading("Decisions Made", level=2)
        h.runs[0].font.name = 'Arial'
        h.runs[0].font.size = Pt(14)
        h.runs[0].font.bold = True

        for item in mom.get("decisions", []):
            p = doc.add_paragraph(item, style='List Bullet')
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(11)
        if not mom.get("decisions"):
            p = doc.add_paragraph("No decisions recorded")
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(11)

        h = doc.add_heading("Action Items", level=2)
        h.runs[0].font.name = 'Arial'
        h.runs[0].font.size = Pt(14)
        h.runs[0].font.bold = True

        action_items = mom.get("action_items", [])
        if action_items:
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Light Shading Accent 1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Task'
            hdr_cells[1].text = 'Owner'
            hdr_cells[2].text = 'Deadline'

            for cell in hdr_cells:
                for p_cell in cell.paragraphs:
                    for r in p_cell.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(10)
                        r.font.bold = True

            for item in action_items:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get("task", "")
                row_cells[1].text = item.get("owner", "Unassigned")
                row_cells[2].text = item.get("deadline", "N/A") or "N/A"
                for cell in row_cells:
                    for p_cell in cell.paragraphs:
                        for r in p_cell.runs:
                            r.font.name = 'Arial'
                            r.font.size = Pt(10)
        else:
            p = doc.add_paragraph("No action items recorded")
            p.runs[0].font.name = 'Arial'
            p.runs[0].font.size = Pt(11)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        filename = f"{title.lower().replace(' ', '_')}_mom.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as exc:
        logger.exception("Failed to generate DOCX")
        raise HTTPException(status_code=500, detail=f"DOCX generation failed: {exc}")


@router.delete("/{meeting_id}")
async def delete_meeting(
    meeting_id: UUID,
    current_user = Depends(get_current_user),
):
    """Delete a meeting verifying ownership, decrementing profile usage safely."""
    settings = get_settings()
    resolved_user_id = str(current_user.id)
    if settings.mock_mode:
        row = None
        for m in MOCK_MEETINGS:
            if str(m["id"]) == str(meeting_id):
                row = m
                break
        if not row:
            raise HTTPException(status_code=404, detail="Meeting not found")
        if str(row["user_id"]) != resolved_user_id:
            raise HTTPException(status_code=403, detail="You do not have permission to access this meeting")
        
        MOCK_MEETINGS.remove(row)
        
        from routes.auth import MOCK_PROFILES
        if resolved_user_id in MOCK_PROFILES:
            MOCK_PROFILES[resolved_user_id]["meetings_used"] = max(0, MOCK_PROFILES[resolved_user_id]["meetings_used"] - 1)
        return {"status": "deleted"}

    try:
        supabase = get_supabase()
    except SupabaseNotConfiguredError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    # Fetch meeting to verify owner
    result = await run_db_query(
        supabase.table("meetings").select("user_id").eq("id", str(meeting_id)).maybe_single().execute
    )
    row = result.data
    if not row:
        raise HTTPException(status_code=404, detail="Meeting not found")
    if str(row["user_id"]) != resolved_user_id:
        raise HTTPException(status_code=403, detail="You do not have permission to access this meeting")

    # Delete from meetings
    await run_db_query(
        supabase.table("meetings").delete().eq("id", str(meeting_id)).execute
    )

    # Decrement meetings_used safely (using MAX(meetings_used - 1, 0))
    profile_result = await run_db_query(
        supabase.table("profiles").select("meetings_used").eq("id", resolved_user_id).maybe_single().execute
    )
    if profile_result.data:
        curr_used = profile_result.data.get("meetings_used", 0)
        await run_db_query(
            supabase.table("profiles").update({"meetings_used": max(0, curr_used - 1)}).eq("id", resolved_user_id).execute
        )

    return {"status": "deleted"}
